from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from types import SimpleNamespace
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
import pytest

from core.models import ChangeStatistics, ComparisonResult, ParsedDocument
from reporters import docx_reporter
from reporters.docx_reporter import DocxTrackChangesReporter


def make_result(tmp_path: Path) -> ComparisonResult:
    file_a = tmp_path / "a.docx"
    file_b = tmp_path / "b.docx"
    file_a.write_text("a", encoding="utf-8")
    file_b.write_text("b", encoding="utf-8")
    doc_a = ParsedDocument([], "DOCX", str(file_a))
    doc_b = ParsedDocument([], "DOCX", str(file_b))
    return ComparisonResult(
        file_a=doc_a,
        file_b=doc_b,
        changes=[],
        statistics=ChangeStatistics.from_changes([]),
        timestamp=datetime.now(timezone.utc),
    )


class FakeDoc:
    def __init__(self, read_only: bool = False) -> None:
        self.closed = False
        self.saved_args = None
        self.ProtectionType = -1  # wdNoProtection
        self.ReadOnly = read_only
        self.Final = False
        self.ReadOnlyRecommended = False

    def Close(self, SaveChanges=0) -> None:
        self.closed = True

    def SaveAs2(self, path, FileFormat=None) -> None:
        self.saved_args = (path, FileFormat)

    def Unprotect(self, password="") -> None:
        pass


class FakeDocuments:
    def __init__(self) -> None:
        self.opened = []

    def Open(self, path, ReadOnly=False, AddToRecentFiles=True):
        self.opened.append((path, ReadOnly))
        return FakeDoc()


class FakeWord:
    def __init__(self) -> None:
        self.Visible = None
        self.DisplayAlerts = None
        self.Documents = FakeDocuments()
        self.Application = self
        self.ActiveDocument = None
        self.compare_kwargs = None
        self.quit_called = False
        self.Version = "16.0"

    def CompareDocuments(self, **kwargs) -> None:
        self.compare_kwargs = kwargs
        self.ActiveDocument = FakeDoc()

    def Quit(self) -> None:
        self.quit_called = True


class ReadOnlyThenWritableDocuments:
    def __init__(self) -> None:
        self.opened = []
        self.docs = [FakeDoc(read_only=True), FakeDoc(read_only=False)]

    def Open(self, path, **kwargs):
        self.opened.append((path, kwargs))
        return self.docs.pop(0)


class FakeFindReplacement:
    def __init__(self) -> None:
        self.Text = ""

    def ClearFormatting(self) -> None:
        pass


class FakeFind:
    def __init__(self, owner_range) -> None:
        self.owner_range = owner_range
        self.Text = ""
        self.Replacement = FakeFindReplacement()

    def ClearFormatting(self) -> None:
        pass

    def Execute(self, *args, **kwargs) -> None:
        replace_mode = kwargs.get("Replace")
        if replace_mode is None and len(args) >= 11:
            replace_mode = args[10]
        if replace_mode != 2:
            return
        self.owner_range.text = self.owner_range.text.replace(self.Text, self.Replacement.Text)


class FakeStoryRange:
    def __init__(self, text: str) -> None:
        self.text = text
        self.NextStoryRange = None
        self.Find = FakeFind(self)


class FakeDocWithStoryRanges:
    def __init__(self, text: str) -> None:
        self._story = FakeStoryRange(text)
        self.StoryRanges = self._story

    @property
    def text(self) -> str:
        return self._story.text


def test_docx_reporter_compare_documents_called(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    result = make_result(tmp_path)
    fake_word = FakeWord()

    def fake_dispatch(_):
        return fake_word

    monkeypatch.setattr(
        docx_reporter, "win32com", SimpleNamespace(client=SimpleNamespace(Dispatch=fake_dispatch))
    )
    monkeypatch.setattr(
        docx_reporter, "pythoncom", SimpleNamespace(CoInitialize=lambda: None, CoUninitialize=lambda: None)
    )

    reporter = DocxTrackChangesReporter(author="Tester")
    output = reporter.generate(result, str(tmp_path / "out.docx"))
    assert Path(output).suffix == ".docx"
    assert fake_word.compare_kwargs is not None
    assert fake_word.compare_kwargs["Destination"] == 2
    assert fake_word.compare_kwargs["Granularity"] == 1
    assert fake_word.compare_kwargs["CompareFormatting"] is True
    assert fake_word.compare_kwargs["CompareCaseChanges"] is True
    assert fake_word.compare_kwargs["CompareWhitespace"] is True
    assert fake_word.compare_kwargs["CompareTables"] is True
    assert fake_word.compare_kwargs["CompareHeaders"] is True
    assert fake_word.compare_kwargs["CompareFootnotes"] is True
    assert fake_word.compare_kwargs["CompareTextboxes"] is True
    assert fake_word.compare_kwargs["CompareFields"] is True
    assert fake_word.compare_kwargs["CompareComments"] is True
    assert fake_word.compare_kwargs["RevisedAuthor"] == "Tester"
    assert fake_word.ActiveDocument is not None
    assert fake_word.ActiveDocument.saved_args is not None
    assert fake_word.ActiveDocument.saved_args[1] == 12
    assert fake_word.quit_called is True


def test_docx_reporter_falls_back_on_error(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    result = make_result(tmp_path)
    fake_word = FakeWord()

    def fake_dispatch(_):
        return fake_word

    def raise_compare(**_kwargs):
        raise RuntimeError("boom")

    fake_word.CompareDocuments = raise_compare

    monkeypatch.setattr(
        docx_reporter, "win32com", SimpleNamespace(client=SimpleNamespace(Dispatch=fake_dispatch))
    )
    monkeypatch.setattr(
        docx_reporter, "pythoncom", SimpleNamespace(CoInitialize=lambda: None, CoUninitialize=lambda: None)
    )

    reporter = DocxTrackChangesReporter()
    output = reporter.generate(result, str(tmp_path / "out.docx"))
    assert Path(output).suffix == ".html"
    assert Path(output).exists()
    assert fake_word.quit_called is True


def test_docx_reporter_is_available_true_false(monkeypatch: pytest.MonkeyPatch) -> None:
    fake_word = FakeWord()

    def fake_dispatch(_):
        return fake_word

    monkeypatch.setattr(
        docx_reporter, "win32com", SimpleNamespace(client=SimpleNamespace(Dispatch=fake_dispatch))
    )
    monkeypatch.setattr(
        docx_reporter, "pythoncom", SimpleNamespace(CoInitialize=lambda: None, CoUninitialize=lambda: None)
    )

    reporter = DocxTrackChangesReporter()
    assert reporter.is_available() is True

    def raise_dispatch(_):
        raise RuntimeError("no word")

    monkeypatch.setattr(
        docx_reporter, "win32com", SimpleNamespace(client=SimpleNamespace(Dispatch=raise_dispatch))
    )
    reporter.startup_timeout = 0.0
    assert reporter.is_available() is False


def test_docx_reporter_fallback_when_unavailable(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    result = make_result(tmp_path)
    reporter = DocxTrackChangesReporter()
    monkeypatch.setattr(docx_reporter, "win32com", None)

    output = reporter.generate(result, str(tmp_path / "report.docx"))
    assert Path(output).suffix == ".html"
    assert Path(output).exists()
    assert not (tmp_path / "report.xlsx").exists()


def test_docx_reporter_reopens_read_only_doc_as_editable(tmp_path: Path) -> None:
    fake_word = FakeWord()
    fake_word.Documents = ReadOnlyThenWritableDocuments()
    reporter = DocxTrackChangesReporter()

    opened_doc = reporter._open_document_for_compare(
        word=fake_word,
        file_path="C:\\temp\\input.docx",
        tmp_dir=str(tmp_path),
        prefix="a_",
    )

    assert opened_doc.ReadOnly is False
    assert len(fake_word.Documents.opened) == 2


def test_docx_reporter_strips_docx_read_only_flags(tmp_path: Path) -> None:
    docx_path = tmp_path / "protected.docx"
    document = Document()
    document.add_paragraph("hello")
    document.save(docx_path)

    settings_member = "word/settings.xml"
    word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    with zipfile.ZipFile(docx_path, "r") as source:
        settings_xml = source.read(settings_member)
    root = ET.fromstring(settings_xml)
    ET.SubElement(
        root,
        f"{{{word_ns}}}documentProtection",
        {
            f"{{{word_ns}}}edit": "readOnly",
            f"{{{word_ns}}}enforcement": "1",
        },
    )
    ET.SubElement(root, f"{{{word_ns}}}readOnlyRecommended")
    patched_settings = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    temp_archive = tmp_path / "protected_tmp.docx"
    with zipfile.ZipFile(docx_path, "r") as source:
        with zipfile.ZipFile(temp_archive, "w") as target:
            for info in source.infolist():
                payload = (
                    patched_settings
                    if info.filename == settings_member
                    else source.read(info.filename)
                )
                target.writestr(info, payload)
    temp_archive.replace(docx_path)

    DocxTrackChangesReporter._strip_docx_protection_flags(str(docx_path))

    with zipfile.ZipFile(docx_path, "r") as source:
        cleaned_settings = ET.fromstring(source.read(settings_member))

    assert cleaned_settings.find(f"{{{word_ns}}}documentProtection") is None
    assert cleaned_settings.find(f"{{{word_ns}}}readOnlyRecommended") is None


def test_docx_reporter_decodes_common_html_entities_in_document() -> None:
    reporter = DocxTrackChangesReporter()
    doc = FakeDocWithStoryRanges("Don&#39;t &amp;#39; &apos; &#x27;")

    reporter._decode_common_html_entities_in_document(doc)

    assert doc.text == "Don't ' ' '"
