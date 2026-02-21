from __future__ import annotations

from pathlib import Path

from docx import Document

from core.models import ParsedDocument, ParseError, Segment, SegmentContext
from parsers.base import BaseParser

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _local(tag: str) -> str:
    """Return local name of an XML tag, stripping namespace."""
    return tag.split("}")[1] if "}" in tag else tag


def _para_style(para_elem) -> str:
    """Extract paragraph style name from a <w:p> element."""
    ppr = para_elem.find(f"{{{_W}}}pPr")
    if ppr is None:
        return ""
    pstyle = ppr.find(f"{{{_W}}}pStyle")
    if pstyle is None:
        return ""
    return pstyle.get(f"{{{_W}}}val", "")


class DocxParser(BaseParser):
    name = "DOCX Parser"
    supported_extensions = [".docx"]
    format_description = "Word Document"

    def can_handle(self, filepath: str) -> bool:
        return Path(filepath).suffix.lower() == ".docx"

    def parse(self, filepath: str) -> ParsedDocument:
        ext = Path(filepath).suffix.lower()
        if ext == ".doc":
            raise NotImplementedError("DOC format requires conversion to DOCX first")

        try:
            document = Document(filepath)
        except Exception as exc:
            raise ParseError(filepath, str(exc)) from exc

        segments: list[Segment] = []
        for seg_id, location, para_elem, style_name in self._iter_paragraphs(document):
            text = self._extract_text(para_elem)
            if not text or not text.strip():
                continue
            context = SegmentContext(
                file_path=filepath,
                location=location,
                position=len(segments) + 1,
                group=None,
            )
            segments.append(
                Segment(
                    id=seg_id,
                    source=None,
                    target=text,
                    context=context,
                    metadata={"style": style_name},
                )
            )

        return ParsedDocument(
            segments=segments,
            format_name=self.format_description,
            file_path=filepath,
            metadata={},
            encoding=None,
        )

    # ------------------------------------------------------------------
    # Paragraph iteration — body, tables, headers, footers
    # ------------------------------------------------------------------

    def _iter_paragraphs(self, document):
        """Yield (seg_id, location, para_element, style) for all paragraphs
        in document order: body, headers/footers, text boxes, footnotes/endnotes."""
        yield from self._walk_element(document.element.body, "body_p", "")
        yield from self._iter_headers_footers(document)
        yield from self._iter_text_boxes(document)
        yield from self._iter_footnotes_endnotes(document)

    def _walk_element(self, element, para_id_prefix, table_id_prefix):
        """Walk any container element (<w:body>, <w:hdr>, <w:ftr>),
        yielding paragraphs and recursing into tables in document order."""
        para_n = 0
        table_n = 0
        for child in element:
            local = _local(child.tag)
            if local == "p":
                para_n += 1
                yield (
                    f"{para_id_prefix}{para_n}",
                    f"Paragraph {para_n}",
                    child,
                    _para_style(child),
                )
            elif local == "tbl":
                table_n += 1
                yield from self._walk_table(child, table_n, id_prefix=table_id_prefix)

    def _walk_table(self, tbl, t_num, id_prefix=""):
        """Recursively walk a <w:tbl>, yielding paragraphs from every cell.
        Handles nested tables via id_prefix to keep IDs unique."""
        row_n = 0
        for child in tbl:
            if _local(child.tag) != "tr":
                continue
            row_n += 1
            col_n = 0
            for cell_child in child:
                if _local(cell_child.tag) != "tc":
                    continue
                col_n += 1
                para_n = 0
                nested_t = 0
                for item in cell_child:
                    local = _local(item.tag)
                    if local == "p":
                        para_n += 1
                        seg_id = f"{id_prefix}t{t_num}_r{row_n}_c{col_n}_p{para_n}"
                        location = (
                            f"Table {t_num}, Row {row_n}, Cell {col_n}, Para {para_n}"
                        )
                        yield seg_id, location, item, _para_style(item)
                    elif local == "tbl":
                        nested_t += 1
                        nested_prefix = f"{id_prefix}t{t_num}_r{row_n}_c{col_n}_"
                        yield from self._walk_table(item, nested_t, id_prefix=nested_prefix)

    def _iter_footnotes_endnotes(self, document):
        """Yield paragraphs from footnotes and endnotes.
        Accesses the footnotes/endnotes parts via document relationships.
        Skips separator notes (id=-1, id=0) which are Word internals."""
        for rel in document.part.rels.values():
            reltype = rel.reltype
            if "/footnotes" in reltype:
                elem = self._get_part_element(rel.target_part)
                if elem is not None:
                    yield from self._walk_notes(elem, "fn", "footnote")
            elif "/endnotes" in reltype:
                elem = self._get_part_element(rel.target_part)
                if elem is not None:
                    yield from self._walk_notes(elem, "en", "endnote")

    @staticmethod
    def _get_part_element(part):
        """Return the lxml element for a part.
        XmlPart exposes .element directly; base Part stores raw bytes in .blob."""
        if hasattr(part, "element"):
            return part.element
        if hasattr(part, "blob") and part.blob:
            from lxml import etree
            try:
                return etree.fromstring(part.blob)
            except Exception:
                return None
        return None

    def _walk_notes(self, notes_elem, id_prefix, note_local):
        """Walk a <w:footnotes> or <w:endnotes> element, yielding paragraphs
        from each individual note. Separator notes (id=-1, id=0) are skipped."""
        id_attr = f"{{{_W}}}id"
        for note in notes_elem:
            if _local(note.tag) != note_local:
                continue
            note_id = note.get(id_attr, "")
            if note_id in ("-1", "0"):
                continue
            yield from self._walk_element(
                note,
                para_id_prefix=f"{id_prefix}{note_id}_p",
                table_id_prefix=f"{id_prefix}{note_id}_",
            )

    def _iter_text_boxes(self, document):
        """Find every <w:txbxContent> in the document and yield its paragraphs.
        Text boxes can live anywhere — body, table cells, headers, footers."""
        txbx_tag = f"{{{_W}}}txbxContent"
        txbx_n = 0
        for txbx in document.element.iter(txbx_tag):
            txbx_n += 1
            yield from self._walk_element(
                txbx,
                para_id_prefix=f"txbx{txbx_n}_p",
                table_id_prefix=f"txbx{txbx_n}_",
            )

    def _iter_headers_footers(self, document):
        """Yield paragraphs from headers and footers of all sections.
        Skips headers/footers that are linked to the previous section
        to avoid emitting the same content multiple times."""
        seen: set[int] = set()
        for sec_n, section in enumerate(document.sections, start=1):
            for hf, label, id_prefix in (
                (section.header, "Header", f"hdr_s{sec_n}_"),
                (section.footer, "Footer", f"ftr_s{sec_n}_"),
            ):
                if hf.is_linked_to_previous:
                    continue
                elem_key = id(hf._element)
                if elem_key in seen:
                    continue
                seen.add(elem_key)
                yield from self._walk_element(
                    hf._element,
                    para_id_prefix=f"{id_prefix}p",
                    table_id_prefix=id_prefix,
                )

    # ------------------------------------------------------------------
    # Text extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_text(para_elem) -> str:
        """Extract plain text from a <w:p> element, preserving tabs and line breaks.
        Does NOT descend into <w:txbxContent> — text boxes are separate segments."""
        text_parts: list[str] = []
        DocxParser._collect_text(para_elem, text_parts)
        return "".join(text_parts)

    @staticmethod
    def _collect_text(elem, parts: list[str]) -> None:
        """Recursive text collector. Stops at <w:txbxContent> boundaries so that
        text box content is not mixed into the paragraph that anchors the shape."""
        local = _local(elem.tag)
        if local == "txbxContent":
            return
        if local == "t":
            parts.append(elem.text or "")
        elif local == "tab":
            parts.append("\t")
        elif local in ("br", "cr"):
            parts.append("\n")
        for child in elem:
            DocxParser._collect_text(child, parts)

    @staticmethod
    def _extract_paragraph_text(paragraph) -> str:
        """Backward-compatible wrapper: accepts a python-docx Paragraph object."""
        return DocxParser._extract_text(paragraph._element)

    # ------------------------------------------------------------------

    def validate(self, filepath: str) -> list[str]:
        errors: list[str] = []
        try:
            ext = Path(filepath).suffix.lower()
            if ext == ".doc":
                errors.append("DOC format requires conversion to DOCX first")
                return errors
            _ = Document(filepath)
        except Exception as exc:
            errors.append(str(exc))
        return errors
